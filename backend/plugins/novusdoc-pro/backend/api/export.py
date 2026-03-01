"""
NovusDoc Pro 导出/导入 API handlers (Word/PDF)

依赖：
  - python-docx (Word 导出/导入)
  - weasyprint 或 pdfkit (PDF 导出, 可选)

如果依赖未安装，返回 501 提示安装。
"""

from __future__ import annotations

import html as html_mod
import io
import json
import tempfile

from starlette.responses import Response

from app.core.logging import get_logger

logger = get_logger("plugin.novusdoc-pro.api")


from .utils import resolve_tenant_id, safe_int as _safe_int


def _content_disposition(title: str, ext: str) -> str:
    """Build Content-Disposition header safe for non-ASCII filenames (RFC 5987)."""
    import re
    from urllib.parse import quote
    name = re.sub(r'[\\/:*?"<>|]', "_", (title or "document").strip())[:100] or "document"
    ascii_name = name.encode("ascii", errors="replace").decode("ascii") + ext
    utf8_name = quote(name + ext, safe="")
    return f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{utf8_name}"


async def _get_doc_content(db, tenant_id: int, doc_id: int):
    """获取文档内容（JSON 和纯文本）— 通过 load_plugin_handler 跨插件调用"""
    # 方式 1：通过插件模块加载器调用 novusdoc service
    try:
        from app.plugins.module_loader import load_plugin_handler
        get_document = load_plugin_handler(
            "novusdoc", "services.document_service.get_document",
        )
        if get_document:
            doc = await get_document(db, tenant_id, doc_id)
            if doc:
                return doc.get("title", ""), doc.get("content")
            return None, None
    except Exception as exc:
        logger.warning("export: failed to call novusdoc service: %s", exc)

    # 方式 2：降级直接查询文档表
    from sqlalchemy import text
    result = await db.execute(
        text(
            "SELECT title, content FROM px_novusdoc_documents "
            "WHERE id = :did AND tenant_id = :tid AND is_deleted = false"
        ),
        {"did": doc_id, "tid": tenant_id},
    )
    row = result.first()
    if not row:
        return None, None
    return row[0], row[1]


def _json_to_text(content) -> str:
    """从 Tiptap JSON 提取纯文本"""
    if not content:
        return ""
    if isinstance(content, str):
        try:
            content = json.loads(content)
        except (json.JSONDecodeError, TypeError):
            return content

    parts = []

    def _walk(node):
        if isinstance(node, dict):
            if node.get("type") == "text":
                parts.append(node.get("text", ""))
            for child in node.get("content", []):
                _walk(child)
            if node.get("type") in ("paragraph", "heading", "blockquote"):
                parts.append("\n")
        elif isinstance(node, list):
            for item in node:
                _walk(item)

    _walk(content)
    return "".join(parts).strip()


async def export_word(request, db, ctx):
    """POST /docs/{doc_id}/export/word"""
    from ..services.license_gate import check_license_valid, license_required_error
    is_valid, license_info = await check_license_valid(ctx)
    if not is_valid:
        return license_required_error(license_info)

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
    except ImportError:
        return {
            "error": "python-docx is not installed. Run: pip install python-docx",
            "code": 5000,
            "status_code": 501,
        }

    tenant_id = resolve_tenant_id(ctx)
    if tenant_id is None:
        return {"error": "tenant_id required", "code": 4010, "status_code": 401}
    doc_id, err = _safe_int(request.path_params.get("doc_id"), "doc_id")
    if err:
        return err

    title, content = await _get_doc_content(db, tenant_id, doc_id)
    if title is None:
        return {"error": "document not found", "code": 4040, "status_code": 404}

    text = _json_to_text(content)

    docx = DocxDocument()
    docx.add_heading(title or "Untitled", level=1)
    for para in text.split("\n"):
        if para.strip():
            docx.add_paragraph(para)

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)

    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _content_disposition(title, ".docx")},
    )


async def export_pdf(request, db, ctx):
    """POST /docs/{doc_id}/export/pdf"""
    from ..services.license_gate import check_license_valid, license_required_error
    is_valid, license_info = await check_license_valid(ctx)
    if not is_valid:
        return license_required_error(license_info)

    tenant_id = resolve_tenant_id(ctx)
    if tenant_id is None:
        return {"error": "tenant_id required", "code": 4010, "status_code": 401}
    doc_id, err = _safe_int(request.path_params.get("doc_id"), "doc_id")
    if err:
        return err

    title, content = await _get_doc_content(db, tenant_id, doc_id)
    if title is None:
        return {"error": "document not found", "code": 4040, "status_code": 404}

    text = _json_to_text(content)

    # 构建简单 HTML
    safe_title = html_mod.escape(title or 'Untitled')
    html_content = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{safe_title}</title>
<style>body {{ font-family: sans-serif; max-width: 800px; margin: 40px auto; line-height: 1.6; }}
h1 {{ border-bottom: 2px solid #333; padding-bottom: 8px; }}</style></head>
<body><h1>{safe_title}</h1>"""
    for para in text.split("\n"):
        if para.strip():
            html_content += f"<p>{html_mod.escape(para)}</p>"
    html_content += "</body></html>"

    # 尝试 weasyprint
    try:
        from weasyprint import HTML as WeasyprintHTML
        pdf_bytes = WeasyprintHTML(string=html_content).write_pdf()
    except ImportError:
        return {
            "error": "weasyprint is not installed. Install weasyprint for PDF support.",
            "code": 5001,
            "status_code": 501,
        }

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": _content_disposition(title, ".pdf")},
    )


async def import_word(request, db, ctx):
    """POST /docs/import/word — 上传 .docx 文件并创建文档"""
    from ..services.license_gate import check_license_valid, license_required_error
    is_valid, license_info = await check_license_valid(ctx)
    if not is_valid:
        return license_required_error(license_info)

    try:
        import mammoth
    except ImportError:
        return {
            "error": "mammoth is not installed. Run: pip install mammoth",
            "code": 5000,
            "status_code": 501,
        }

    tenant_id = resolve_tenant_id(ctx)
    if tenant_id is None:
        return {"error": "tenant_id required", "code": 4010, "status_code": 401}

    form = await request.form()
    file = form.get("file")
    if not file:
        return {"error": "file required", "code": 4001, "status_code": 400}

    file_bytes = await file.read()

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        tmp.seek(0)
        result = mammoth.convert_to_html(tmp)
        html = result.value

    # 从 HTML 提取标题（第一个 h1 或文件名）
    import re
    title_match = re.search(r"<h1[^>]*>(.*?)</h1>", html, re.IGNORECASE)
    title = title_match.group(1) if title_match else (file.filename or "Imported Document").replace(".docx", "")

    return {
        "title": title,
        "html": html,
        "message": "Word file converted. Use the HTML content to create a document.",
    }
