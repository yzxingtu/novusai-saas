"""
Document Parser Module / 文档解析器模块

Supports 11 formats: PDF, DOCX, TXT, Markdown, CSV, XLSX, HTML, Q&A, URL, PPTX, Image (JPG/PNG/WebP/GIF).
Unified output as ParsedPage list for chunker consumption.
支持 PDF、DOCX、TXT、Markdown、CSV、XLSX、HTML、Q&A、URL、PPTX、图片（JPG/PNG/WebP/GIF）十一种格式解析，
统一输出 ParsedPage 列表供分块器使用。
"""

from __future__ import annotations

from typing import TYPE_CHECKING, BinaryIO

if TYPE_CHECKING:
    from app.ai.rag.audio_describer import AudioDescriber
    from app.ai.rag.video_describer import VideoDescriber
    from app.ai.rag.vision_describer import VisionDescriber
    from app.models.ai.knowledge_base import KnowledgeBase

from app.ai.rag.html_section_support import build_html_sections
from app.ai.rag.parser_contracts import DocumentParser, ParsedPage, QaPairParser
from app.ai.rag.parser_multimodal_support import (
    AUDIO_TYPES as _AUDIO_TYPES,
)
from app.ai.rag.parser_multimodal_support import (
    IMAGE_TYPES as _IMAGE_TYPES,
)
from app.ai.rag.parser_multimodal_support import (
    VIDEO_TYPES as _VIDEO_TYPES,
)
from app.ai.rag.parser_multimodal_support import (
    AudioParser,
    ImageParser,
    VideoParser,
)
from app.ai.rag.parser_pptx_support import PptxParser
from app.ai.rag.url_fetcher import fetch_public_url_text
from app.ai.text_semantics import parse_markdown_heading, split_on_blank_lines
from app.core.i18n import _
from app.core.logging import LogManager
from app.exceptions import BusinessException

logger = LogManager.get_logger("ai.rag.parser")


class PdfParser(DocumentParser):
    """
    PDF Parser / PDF 解析器

    Uses PyMuPDF (fitz) to extract text, preserving page number metadata.
    When knowledge_base.extract_images=True and VisionDescriber is injected,
    also extracts embedded images and calls Vision model to generate descriptions.
    使用 PyMuPDF (fitz) 提取文本，保留页码元数据。
    当 knowledge_base.extract_images=True 且注入了 VisionDescriber 时，
    同时提取内嵌图片并调用 Vision 模型生成图片描述追加到页面列表。

    Image filter: images smaller than 4KB (noise/decoration) are skipped.
    图片过滤规则：小于 4KB（噪点/装饰图）的图片跳过。
    """

    # Skip images below this size (avoid processing logos, separators, etc.) / 跳过小于此大小的图片（避免处理 logo、分隔线等装饰图片）
    _MIN_IMAGE_BYTES = 4 * 1024

    def __init__(
        self,
        vision_describer: VisionDescriber | None = None,
        knowledge_base: KnowledgeBase | None = None,
    ) -> None:
        self._vision_describer = vision_describer
        self._knowledge_base = knowledge_base

    async def parse(
        self, file_content: BinaryIO, file_name: str = ""
    ) -> list[ParsedPage]:
        import fitz  # PyMuPDF / PDF 解析库

        pages: list[ParsedPage] = []
        data = file_content.read()
        doc = fitz.open(stream=data, filetype="pdf")

        extract_images: bool = (
            self._vision_describer is not None
            and self._knowledge_base is not None
            and getattr(self._knowledge_base, "extract_images", False)
        )

        try:
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text").strip()
                if text:
                    pages.append(
                        ParsedPage(
                            content=text,
                            metadata={
                                "page": page_num + 1,
                                "source": file_name,
                            },
                        )
                    )

                if not extract_images:
                    continue

                for img_idx, img_info in enumerate(page.get_images(full=True)):
                    xref = img_info[0]
                    try:
                        base_image = doc.extract_image(xref)
                    except Exception:
                        continue
                    if not base_image:
                        continue

                    image_bytes: bytes = base_image.get("image", b"")
                    if not image_bytes or len(image_bytes) < self._MIN_IMAGE_BYTES:
                        continue

                    ext: str = base_image.get("ext", "png")
                    mime_type = (
                        "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
                    )

                    description = await self._vision_describer.describe_image(  # type: ignore[union-attr]  # 可选依赖 / optional
                        image_bytes=image_bytes,
                        mime_type=mime_type,
                        knowledge_base=self._knowledge_base,  # type: ignore[arg-type]  # 可选依赖 / optional
                    )
                    if description:
                        pages.append(
                            ParsedPage(
                                content=description,
                                metadata={
                                    "page": page_num + 1,
                                    "image_index": img_idx,
                                    "source": file_name,
                                    "type": "image",
                                },
                            )
                        )
        finally:
            doc.close()

        logger.info(
            "PDF parsed: {}, pages={} (extract_images={})",
            file_name,
            len(pages),
            extract_images,
        )
        return pages


class DocxParser(DocumentParser):
    """
    DOCX Parser / DOCX 解析器

    Uses python-docx to extract paragraphs and tables, preserving heading hierarchy metadata.
    使用 python-docx 提取段落和表格，保留标题层级元数据。
    """

    async def parse(
        self, file_content: BinaryIO, file_name: str = ""
    ) -> list[ParsedPage]:
        from docx import Document

        doc = Document(file_content)
        pages: list[ParsedPage] = []
        current_heading: str | None = None
        paragraph_index = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect heading / 检测标题
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                current_heading = text

            pages.append(
                ParsedPage(
                    content=text,
                    metadata={
                        "heading": current_heading,
                        "paragraph": paragraph_index,
                        "source": file_name,
                    },
                )
            )
            paragraph_index += 1

        # Process tables: convert to Markdown format / 处理表格：转为 Markdown 格式
        for table_idx, table in enumerate(doc.tables):
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append("| " + " | ".join(cells) + " |")
            if rows_text:
                # Insert separator row / 插入分隔符行
                header_sep = (
                    "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
                )
                table_md = (
                    rows_text[0] + "\n" + header_sep + "\n" + "\n".join(rows_text[1:])
                )
                pages.append(
                    ParsedPage(
                        content=table_md,
                        metadata={
                            "table_index": table_idx,
                            "source": file_name,
                        },
                    )
                )

        logger.info(
            "DOCX parsed: {}, segments={}",
            file_name,
            len(pages),
        )
        return pages


class TxtParser(DocumentParser):
    """
    TXT Parser / TXT 解析器

    Splits by blank lines into paragraphs.
    按空行分段。
    """

    async def parse(
        self, file_content: BinaryIO, file_name: str = ""
    ) -> list[ParsedPage]:
        text = file_content.read().decode("utf-8", errors="replace")
        paragraphs = split_on_blank_lines(text)

        pages: list[ParsedPage] = []
        for idx, para in enumerate(paragraphs):
            content = para.strip()
            if content:
                pages.append(
                    ParsedPage(
                        content=content,
                        metadata={
                            "paragraph": idx,
                            "source": file_name,
                        },
                    )
                )

        logger.info(
            "TXT parsed: {}, paragraphs={}",
            file_name,
            len(pages),
        )
        return pages


class MarkdownParser(DocumentParser):
    """
    Markdown Parser / Markdown 解析器

    Splits by heading hierarchy, preserving headings as metadata.
    Code blocks and tables are kept intact.
    按标题层级分段，保留标题作为元数据。
    代码块和表格完整保留。
    """

    async def parse(
        self, file_content: BinaryIO, file_name: str = ""
    ) -> list[ParsedPage]:
        text = file_content.read().decode("utf-8", errors="replace")
        pages: list[ParsedPage] = []
        current_heading: str | None = None
        current_content: list[str] = []

        def flush():
            """Flush accumulated content as a ParsedPage / 将当前累积内容输出为一个 ParsedPage"""
            nonlocal current_content
            content = "\n".join(current_content).strip()
            if content:
                pages.append(
                    ParsedPage(
                        content=content,
                        metadata={
                            "heading": current_heading,
                            "source": file_name,
                        },
                    )
                )
            current_content = []

        for line in text.split("\n"):
            # Detect heading line / 检测标题行
            heading = parse_markdown_heading(line)
            if heading is not None:
                flush()
                current_heading = heading[1]
                current_content.append(line)
            else:
                current_content.append(line)

        flush()

        logger.info(
            "Markdown parsed: {}, sections={}",
            file_name,
            len(pages),
        )
        return pages


class CsvParser(DocumentParser):
    """
    CSV Parser / CSV 解析器

    Uses pandas to read, each row converted to "column_name: value" format text.
    使用 pandas 读取，每行转为 "列名: 值" 格式文本。
    """

    async def parse(
        self, file_content: BinaryIO, file_name: str = ""
    ) -> list[ParsedPage]:
        import pandas as pd

        df = pd.read_csv(file_content, dtype=str)
        pages: list[ParsedPage] = []
        columns = list(df.columns)

        for row_idx, row in df.iterrows():
            parts = []
            for col in columns:
                val = row.get(col, "")
                if pd.notna(val) and str(val).strip():
                    parts.append(f"{col}: {val}")
            if parts:
                pages.append(
                    ParsedPage(
                        content="\n".join(parts),
                        metadata={
                            "row_index": int(row_idx),
                            "columns": columns,
                            "source": file_name,
                        },
                    )
                )

        logger.info(
            "CSV parsed: {}, rows={}",
            file_name,
            len(pages),
        )
        return pages


class XlsxParser(DocumentParser):
    """
    Excel (.xlsx) Parser / Excel (.xlsx) 解析器

    Uses openpyxl to read, supports multiple sheets, each row converted to "column_name: value" format text.
    使用 openpyxl 读取，支持多 Sheet，每行转为 "列名: 值" 格式文本。
    """

    async def parse(
        self, file_content: BinaryIO, file_name: str = ""
    ) -> list[ParsedPage]:
        from openpyxl import load_workbook

        wb = load_workbook(file_content, read_only=True, data_only=True)
        pages: list[ParsedPage] = []

        try:
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue

                # First row as column names / 第一行作为列名
                headers = [
                    str(h).strip() if h is not None else f"col_{i}"
                    for i, h in enumerate(rows[0])
                ]

                for row_idx, row in enumerate(rows[1:], start=1):
                    parts = []
                    for col_idx, val in enumerate(row):
                        if val is not None and str(val).strip():
                            col_name = (
                                headers[col_idx]
                                if col_idx < len(headers)
                                else f"col_{col_idx}"
                            )
                            parts.append(f"{col_name}: {val}")
                    if parts:
                        pages.append(
                            ParsedPage(
                                content="\n".join(parts),
                                metadata={
                                    "sheet": sheet_name,
                                    "row_index": row_idx,
                                    "columns": headers,
                                    "source": file_name,
                                },
                            )
                        )
        finally:
            wb.close()

        logger.info(
            "XLSX parsed: {}, sheets={}, rows={}",
            file_name,
            len(wb.sheetnames),
            len(pages),
        )
        return pages


class HtmlParser(DocumentParser):
    """
    HTML File Parser / HTML 文件解析器

    Uses BeautifulSoup to extract body content, filtering out script/style and other non-content tags.
    使用 BeautifulSoup 提取正文，过滤 script/style 等非内容标签。
    """

    async def parse(
        self, file_content: BinaryIO, file_name: str = ""
    ) -> list[ParsedPage]:
        html_text = file_content.read().decode("utf-8", errors="replace")
        sections = build_html_sections(
            html_text,
            source=file_name,
        )
        pages = [
            ParsedPage(
                content=section["content"],
                metadata=section["metadata"],
            )
            for section in sections
        ]

        logger.info(
            "HTML parsed: {}, sections={}",
            file_name,
            len(pages),
        )
        return pages


class UrlParser(DocumentParser):
    """
    URL Parser / URL 解析器

    Uses httpx + beautifulsoup4 to fetch web pages and extract body content.
    使用 httpx + beautifulsoup4 爬取网页，提取正文。
    """

    async def parse(
        self, file_content: BinaryIO, file_name: str = ""
    ) -> list[ParsedPage]:
        """
        Parse web page content / 解析网页内容

        file_content stores URL string (UTF-8 encoded).
        file_content 中存储的是 URL 字符串（UTF-8 编码）。
        """
        url = file_content.read().decode("utf-8").strip()
        if not url:
            raise BusinessException(
                message=_("knowledge_base.document.error.parse_failed")
            )

        html_text = await fetch_public_url_text(url)
        sections = build_html_sections(
            html_text,
            source=file_name or url,
            source_url=url,
        )
        pages = [
            ParsedPage(
                content=section["content"],
                metadata=section["metadata"],
            )
            for section in sections
        ]

        logger.info(
            "URL parsed: {}, sections={}",
            url,
            len(pages),
        )
        return pages


def get_parser(
    file_type: str,
    vision_describer: VisionDescriber | None = None,
    audio_describer: AudioDescriber | None = None,
    video_describer: VideoDescriber | None = None,
    knowledge_base: KnowledgeBase | None = None,
) -> DocumentParser:
    """
    Factory method: get parser by file type
    工厂方法：根据文件类型获取解析器

    Args:
        file_type: File type (pdf/docx/txt/md/csv/url/pptx/image/audio/video...) / 文件类型
        vision_describer: Vision description service for ImageParser and enhanced PdfParser
                          图片描述服务，供 ImageParser 和增强版 PdfParser 使用
        audio_describer: Audio-to-text service for AudioParser / 音频转文本服务，供 AudioParser 使用
        video_describer: Video-to-text service for VideoParser / 视频转文本服务，供 VideoParser 使用
        knowledge_base: KB object for image/audio/video parsers / 知识库对象，供图片/音频/视频解析器使用

    Returns:
        Corresponding parser instance / 对应的解析器实例

    Raises:
        BusinessException: Unsupported file type / 不支持的文件类型
    """
    if file_type == "pdf":
        return PdfParser(
            vision_describer=vision_describer,
            knowledge_base=knowledge_base,
        )

    text_parsers: dict[str, type[DocumentParser]] = {
        "docx": DocxParser,
        "txt": TxtParser,
        "md": MarkdownParser,
        "csv": CsvParser,
        "xlsx": XlsxParser,
        "html": HtmlParser,
        "url": UrlParser,
        "pptx": PptxParser,
    }

    if file_type in text_parsers:
        return text_parsers[file_type]()

    if file_type in _IMAGE_TYPES:
        return ImageParser(
            vision_describer=vision_describer,
            knowledge_base=knowledge_base,
        )

    if file_type in _AUDIO_TYPES:
        return AudioParser(
            audio_describer=audio_describer,
            knowledge_base=knowledge_base,
        )

    if file_type in _VIDEO_TYPES:
        return VideoParser(
            video_describer=video_describer,
            knowledge_base=knowledge_base,
        )

    raise BusinessException(
        message=_("knowledge_base.document.error.unsupported_type"),
    )


__all__ = [
    "ParsedPage",
    "DocumentParser",
    "PdfParser",
    "DocxParser",
    "TxtParser",
    "MarkdownParser",
    "CsvParser",
    "XlsxParser",
    "HtmlParser",
    "QaPairParser",
    "UrlParser",
    "PptxParser",
    "ImageParser",
    "AudioParser",
    "VideoParser",
    "get_parser",
]
